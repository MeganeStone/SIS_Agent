# translate_word.py
import os
import gc
from docx import Document
from docx.oxml.ns import qn
from langchain_core.tools import ToolException
from translate_text import translate_text, _translation_cache

# ---------------------- 全局配置 ----------------------
DEFAULT_INPUT_DIR = r"D:\seki\AI\copilotTest\input"
DEFAULT_OUTPUT_DIR = r"D:\seki\AI\copilotTest\output"
DEFAULT_TARGET_LANG = "日语"
DEFAULT_DELAY = 1.2


# ---------- 辅助函数 ----------
def _paragraph_has_picture(paragraph):
    """判断段落是否包含图片（内联图形）"""
    for run in paragraph.runs:
        if run._element.find(qn('w:drawing')) is not None:
            return True
    return False

def _translate_textboxes(doc, target_lang, delay, context):
    """
    翻译文档中所有文本框内的文本。
    使用 local-name() 避免命名空间参数问题。
    """
    try:
        root = doc.part.element
        # 查找所有 w:txbxContent 元素（忽略命名空间）
        txbx_contents = root.xpath('.//*[local-name()="txbxContent"]')
        for txbx in txbx_contents:
            # 遍历文本框内的所有段落
            for para_elem in txbx.xpath('.//*[local-name()="p"]'):
                # 提取所有文本节点内容
                original = ''.join(para_elem.xpath('.//*[local-name()="t"]/text()')).strip()
                if original and len(original) > 1:
                    translated = translate_text(original, target_lang, delay, context=context)
                    if translated:
                        # 替换所有 w:t 节点的文本
                        for t_elem in para_elem.xpath('.//*[local-name()="t"]'):
                            t_elem.text = translated

    except Exception as e:
        print(f"文本框翻译出错: {e}")


# ---------- 主函数 ----------
def translate_word_file(file_name: str, source_dir: str = DEFAULT_INPUT_DIR,
                        output_dir: str = DEFAULT_OUTPUT_DIR,
                        target_lang: str = DEFAULT_TARGET_LANG,
                        delay: float = DEFAULT_DELAY) -> str:
    """
    翻译 Word 文件（仅支持 .docx）
    支持段落、表格单元格、文本框内的文本翻译。
    """
    ext = os.path.splitext(file_name)[1].lower()
    if ext == ".doc":
        raise ToolException("暂不支持 .doc 格式，请将文件转换为 .docx 格式后再试。")
    if ext != ".docx":
        raise ToolException(f"不支持的文件类型: {ext}，仅支持 .docx")

    global _translation_cache
    _translation_cache = {}

    input_path = os.path.abspath(os.path.join(source_dir, file_name))
    output_path = os.path.abspath(os.path.join(
        output_dir,
        f"{os.path.splitext(file_name)[0]}_{target_lang}.docx"
    ))

    if not os.path.exists(input_path):
        raise ToolException(f"Word 文件不存在: {input_path}")

    os.makedirs(output_dir, exist_ok=True)

    doc = Document(input_path)

    # 统计可翻译的文本单元数量（用于进度显示）
    total_units = 0
    for para in doc.paragraphs:
        if not _paragraph_has_picture(para) and para.text.strip():
            total_units += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not _paragraph_has_picture(para) and para.text.strip():
                        total_units += 1

    context = {
        "file_name": file_name,
        "slide_num": 0,
        "total_slides": total_units,
        "translated_segments": [],
        "term_requirements": "TBOX译为TBOX、TSU译为TSU、CAN总线译为CANバス（日语）/CAN Bus（英语），公司名称不翻译"
    }

    print(f"\n开始翻译 Word 文档，共约 {total_units} 个文本单元...")
    processed = 0

    # 1. 翻译主体段落
    for para in doc.paragraphs:
        if _paragraph_has_picture(para):
            continue
        original = para.text.strip()
        if not original or len(original) <= 1:
            continue
        processed += 1
        context["slide_num"] = processed
        print(f"处理段落 {processed}/{total_units}，长度: {len(original)}")
        try:
            translated = translate_text(original, target_lang, delay, context=context)
            if translated:
                para.text = translated
                context["translated_segments"].append(f"原文：{original} | 译文：{translated}")
                if len(context["translated_segments"]) > 5:
                    context["translated_segments"].pop(0)
        except Exception as e:
            print(f"翻译段落失败: {e}")

    # 2. 翻译表格单元格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _paragraph_has_picture(para):
                        continue
                    original = para.text.strip()
                    if not original or len(original) <= 1:
                        continue
                    processed += 1
                    context["slide_num"] = processed
                    print(f"处理单元格段落 {processed}/{total_units}，长度: {len(original)}")
                    try:
                        translated = translate_text(original, target_lang, delay, context=context)
                        if translated:
                            para.text = translated
                            context["translated_segments"].append(f"原文：{original} | 译文：{translated}")
                            if len(context["translated_segments"]) > 5:
                                context["translated_segments"].pop(0)
                    except Exception as e:
                        print(f"翻译单元格段落失败: {e}")

    # 3. 翻译文本框
    print("\n开始翻译文本框内容...")
    _translate_textboxes(doc, target_lang, delay, context)
    print("文本框翻译完成")

    doc.save(output_path)
    print(f"\n✅ Word 翻译完成！输出路径: {output_path}")

    _translation_cache.clear()
    gc.collect()
    return f"Word翻译完成！输出路径: {output_path}"